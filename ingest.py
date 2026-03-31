import os
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangDocument
from pypdf import PdfReader
from docx import Document as DocxDocument
import pandas as pd
from dotenv import load_dotenv

# Load API Keys
load_dotenv()

# --- CONFIGURATION ---
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, "data")
VECTOR_STORE_PATH = os.path.join(BASE_DIR, "faiss_index")
EMBEDDING_MODEL = "text-embedding-3-small"

# Create data directory if it doesn't exist
if not os.path.exists(DATA_DIR):
    os.makedirs(DATA_DIR)
    print(f"Created missing folder: {DATA_DIR}. Drop your files here!")

# --- FILE EXTRACTOR LOGIC ---
def extract_text_from_file(file_path):
    """
    Identifies file type and extracts raw text.
    """
    ext = file_path.split('.')[-1].lower()
    try:
        if ext == 'txt':
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        elif ext == 'pdf':
            reader = PdfReader(file_path)
            return "\n".join([p.extract_text() for p in reader.pages if p.extract_text()])
        elif ext == 'docx':
            doc = DocxDocument(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        elif ext in ['xlsx', 'csv']:
            df = pd.read_excel(file_path) if ext != 'csv' else pd.read_csv(file_path)
            return df.to_string()
    except Exception as e:
        print(f"❌ Error reading {file_path}: {e}")
    return None

# --- MAIN INGESTION ENGINE ---
def ingest_documents():
    embeddings = OpenAIEmbeddings(model=EMBEDDING_MODEL)
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    all_chunks = []

    files = [f for f in os.listdir(DATA_DIR) if os.path.isfile(os.path.join(DATA_DIR, f))]
    
    if not files:
        print(f"No files found in {DATA_DIR}. Please add some files and try again.")
        return

    print(f"📂 Found {len(files)} files. Starting extraction...")

    for fname in files:
        file_path = os.path.join(DATA_DIR, fname)
        raw_text = extract_text_from_file(file_path)
        
        if raw_text:
            chunks = text_splitter.split_text(raw_text)
            for chunk in chunks:
                all_chunks.append(LangDocument(page_content=chunk, metadata={"source": fname}))
            print(f"✅ Processed: {fname} ({len(chunks)} chunks)")

    if not all_chunks:
        print("No valid text extracted. Ingestion aborted.")
        return

    # 3. Create/Save Vector Store
    print(f"🧠 Generating embeddings for {len(all_chunks)} chunks...")
    vector_store = FAISS.from_documents(all_chunks, embeddings)
    vector_store.save_local(VECTOR_STORE_PATH)
    print(f"🚀 SUCCESS! Index saved to '{VECTOR_STORE_PATH}'.")

if __name__ == "__main__":
    ingest_documents()